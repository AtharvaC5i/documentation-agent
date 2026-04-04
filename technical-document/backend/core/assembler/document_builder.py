import os
from pathlib import Path
from datetime import date
from typing import List, Dict, Any
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.parts.numbering import NumberingPart
import re

load_dotenv()

STORAGE_DIR = os.getenv("STORAGE_DIR", os.path.join("..", "storage"))

PAGE_WIDTH_CM  = Cm(21)
PAGE_HEIGHT_CM = Cm(29.7)


# ─────────────────────────────────────────────────────────────
# DESIGN TOKENS
# ─────────────────────────────────────────────────────────────

# ─────────────────────────────────────────────────────────────
# DESIGN TOKENS  (replace the existing class)
# ─────────────────────────────────────────────────────────────

class DesignTokens:
    COLOR_INK          = RGBColor(0x0F, 0x0F, 0x1A)
    COLOR_HEADING      = RGBColor(0x0D, 0x1B, 0x3E)
    COLOR_ACCENT       = RGBColor(0x1A, 0x56, 0xDB)
    COLOR_MUTED        = RGBColor(0x5A, 0x5A, 0x72)
    COLOR_LIGHT_MUTED  = RGBColor(0x9C, 0xA3, 0xAF)
    COLOR_CODE_TEXT    = RGBColor(0x1E, 0x3A, 0x8A)
    COLOR_TABLE_STRIPE = RGBColor(0xF8, 0xF9, 0xFF)
    COLOR_WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
    COLOR_BLOCKQUOTE   = RGBColor(0x2D, 0x4A, 0x8A)
    COLOR_ACCENT_LIGHT = RGBColor(0xEB, 0xF0, 0xFF)   # very light blue for meta bg

    FONT_BODY    = "Calibri"
    FONT_HEADING = "Calibri"
    FONT_MONO    = "Consolas"

    SIZE_DISPLAY   = Pt(42)    # was 32 — too small
    SIZE_SUBTITLE  = Pt(11)
    SIZE_H1        = Pt(16)
    SIZE_H2        = Pt(13)
    SIZE_H3        = Pt(11)
    SIZE_H4        = Pt(10.5)
    SIZE_BODY      = Pt(10.5)
    SIZE_CAPTION   = Pt(9)
    SIZE_CODE      = Pt(9)
    SIZE_META_LABEL = Pt(8)
    SIZE_META_VALUE = Pt(10.5)

    SPACE_BEFORE_H1   = Pt(18)
    SPACE_AFTER_H1    = Pt(6)
    SPACE_BEFORE_H2   = Pt(14)
    SPACE_AFTER_H2    = Pt(4)
    SPACE_BEFORE_H3   = Pt(8)
    SPACE_AFTER_H3    = Pt(3)
    SPACE_AFTER_BODY  = Pt(6)
    LINE_SPACING_BODY = Pt(15)

    MARGIN_TOP    = Cm(2.8)
    MARGIN_BOTTOM = Cm(2.8)
    MARGIN_LEFT   = Cm(3.0)
    MARGIN_RIGHT  = Cm(2.5)


T = DesignTokens

BODY_WIDTH = PAGE_WIDTH_CM - T.MARGIN_LEFT - T.MARGIN_RIGHT  # ~Cm(15.5)


# ─────────────────────────────────────────────────────────────
# LOW-LEVEL XML HELPERS
# ─────────────────────────────────────────────────────────────

def _set_cell_shading(cell, fill_hex: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def _set_paragraph_shading(paragraph, fill_hex: str):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    pPr.append(shd)


def _set_cell_border(cell, sides: dict):
    tc        = cell._tc
    tcPr      = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, attrs in sides.items():
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   attrs.get("val",   "single"))
        el.set(qn("w:sz"),    attrs.get("sz",    "4"))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), attrs.get("color", "000000"))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _remove_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def _set_table_width(table, width_emu: int):
    """Set total table width in EMU via tblW XML."""
    dxa   = int(width_emu * 1440 / 914400)
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"),    str(dxa))
    tblW.set(qn("w:type"), "dxa")


def _set_col_width(table, col_idx: int, width_emu: int):
    dxa = int(width_emu * 1440 / 914400)
    for row in table.rows:
        tc   = row.cells[col_idx]._tc
        tcPr = tc.get_or_add_tcPr()
        tcW  = tcPr.find(qn("w:tcW"))
        if tcW is None:
            tcW = OxmlElement("w:tcW")
            tcPr.append(tcW)
        tcW.set(qn("w:w"),    str(dxa))
        tcW.set(qn("w:type"), "dxa")


def _make_table(container, rows: int, cols: int, total_width_emu: int):
    """
    THE SINGLE CORRECT WAY to add a table in python-docx regardless of container type.

    python-docx's Document.add_table() accepts (rows, cols, style).
    Header/Footer._Body.add_table() is BlockItemContainer.add_table() which
    requires a positional `width` argument.

    To handle BOTH cases uniformly, we build the table XML directly and insert
    it into the container's body element. This bypasses the conflicting signatures.
    """
    from docx.oxml.table import CT_Tbl
    from docx.table import Table

    # Build a minimal table XML element
    tbl_xml = CT_Tbl.new_tbl(rows, cols, Inches(1))  # Inches(1) is a throwaway default
    tbl_obj = Table(tbl_xml, container)

    # Now set the real width we want
    tbl_obj.autofit   = False
    tbl_obj.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_width(tbl_obj, total_width_emu)

    # Insert the tbl element into the container's body
    # Both Document._body and _HeaderFooter._body expose the same _body._body XML element
    try:
        # Document has ._body._body (the actual XML body element)
        body_elem = container._body._body
    except AttributeError:
        try:
            # Header/Footer exposes ._body directly as the XML element
            body_elem = container._body
        except AttributeError:
            body_elem = container._element

    body_elem.append(tbl_xml)
    return tbl_obj


def _prevent_table_row_split(row):
    tr   = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    cantSplit = OxmlElement("w:cantSplit")
    cantSplit.set(qn("w:val"), "1")
    trPr.append(cantSplit)


def _add_horizontal_rule(doc, color_hex: str = "D0D8F0", thickness: str = "4"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm  = OxmlElement("w:bottom")
    btm.set(qn("w:val"),   "single")
    btm.set(qn("w:sz"),    thickness)
    btm.set(qn("w:space"), "1")
    btm.set(qn("w:color"), color_hex)
    pBdr.append(btm)
    pPr.append(pBdr)


def _add_page_number(paragraph):
    run      = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText      = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _add_numpages_field(paragraph):
    run      = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText      = OxmlElement("w:instrText")
    instrText.text = "NUMPAGES"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# ─────────────────────────────────────────────────────────────
# NUMBERING
# ─────────────────────────────────────────────────────────────

def _ensure_numbering_part(doc: Document):
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        numbering_part = NumberingPart.new()
        doc.part.relate_to(
            numbering_part,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering",
        )
    return numbering_part


def _create_fresh_numbering(doc: Document) -> int:
    numbering_part   = _ensure_numbering_part(doc)
    numbering        = numbering_part._element
    abstract_nums    = numbering.findall(qn("w:abstractNum"))
    next_abstract_id = str(len(abstract_nums))

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), next_abstract_id)

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract_num.append(multi_level)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)

    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    pPr_lvl = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"),    "360")
    ind.set(qn("w:hanging"), "360")
    pPr_lvl.append(ind)
    lvl.append(pPr_lvl)
    abstract_num.append(lvl)

    first_num = numbering.find(qn("w:num"))
    if first_num is not None:
        numbering.insert(list(numbering).index(first_num), abstract_num)
    else:
        numbering.append(abstract_num)

    existing_nums = numbering.findall(qn("w:num"))
    next_num_id   = str(len(existing_nums) + 1)

    num_el = OxmlElement("w:num")
    num_el.set(qn("w:numId"), next_num_id)

    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), next_abstract_id)
    num_el.append(abstract_ref)

    numbering.append(num_el)
    return int(next_num_id)


def _apply_num_id_to_paragraph(paragraph, num_id: int, ilvl: int = 0):
    pPr      = paragraph._p.get_or_add_pPr()
    numPr    = OxmlElement("w:numPr")
    ilvl_el  = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    existing = pPr.find(qn("w:numPr"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(numPr)


# ─────────────────────────────────────────────────────────────
# DOCUMENT SETUP
# ─────────────────────────────────────────────────────────────

def _configure_margins(doc: Document):
    for section in doc.sections:
        section.top_margin    = T.MARGIN_TOP
        section.bottom_margin = T.MARGIN_BOTTOM
        section.left_margin   = T.MARGIN_LEFT
        section.right_margin  = T.MARGIN_RIGHT


def _apply_base_styles(doc: Document):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name        = T.FONT_BODY
    normal.font.size        = T.SIZE_BODY
    normal.font.color.rgb   = T.COLOR_INK
    normal.paragraph_format.space_after  = T.SPACE_AFTER_BODY
    normal.paragraph_format.line_spacing = T.LINE_SPACING_BODY

    style_names = [s.name for s in styles]
    for style_name, size, before, after in [
        ("Heading 1", T.SIZE_H1, T.SPACE_BEFORE_H1, T.SPACE_AFTER_H1),
        ("Heading 2", T.SIZE_H2, T.SPACE_BEFORE_H2, T.SPACE_AFTER_H2),
        ("Heading 3", T.SIZE_H3, T.SPACE_BEFORE_H3, T.SPACE_AFTER_H3),
    ]:
        if style_name not in style_names:
            continue
        h = styles[style_name]
        h.font.name       = T.FONT_HEADING
        h.font.size       = size
        h.font.bold       = True
        h.font.color.rgb  = T.COLOR_HEADING if style_name != "Heading 3" else T.COLOR_ACCENT
        h.font.underline  = False
        h.font.italic     = False
        h.paragraph_format.space_before   = before
        h.paragraph_format.space_after    = after
        h.paragraph_format.keep_with_next = True

    try:
        lb = styles["List Bullet"]
        lb.font.name      = T.FONT_BODY
        lb.font.size      = T.SIZE_BODY
        lb.font.color.rgb = T.COLOR_INK
        lb.paragraph_format.space_after = Pt(3)
        lb.paragraph_format.left_indent = Inches(0.25)
    except KeyError:
        pass


# ─────────────────────────────────────────────────────────────
# HEADER & FOOTER  — fixed for Header/Footer containers
# ─────────────────────────────────────────────────────────────

def _add_table_to_hf(hf_container, rows: int, cols: int, total_width_emu: int):
    """
    Add a table to a Header or Footer container.

    python-docx Header/Footer objects inherit add_table() from
    BlockItemContainer which requires a positional `width` argument
    (unlike Document.add_table). We bypass this entirely by building
    the tbl XML directly and appending it to the container's _element.
    """
    from docx.oxml.table import CT_Tbl
    from docx.table import Table

    # Build a skeleton table element (Inches(1) width is thrown away below)
    tbl_xml = CT_Tbl.new_tbl(rows, cols, Inches(1))
    tbl_obj = Table(tbl_xml, hf_container)
    tbl_obj.autofit   = False
    tbl_obj.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Append to the header/footer XML body
    hf_container._element.append(tbl_xml)

    # Now override the width properly
    _set_table_width(tbl_obj, total_width_emu)
    return tbl_obj


def _configure_header_footer(doc: Document, project_name: str):
    # Clear title-page header/footer
    title_section = doc.sections[0]
    title_section.different_first_page_header_footer = False
    title_section.header.is_linked_to_previous = False
    title_section.footer.is_linked_to_previous = False
    for p in title_section.header.paragraphs:
        for r in p.runs: r.text = ""
    for p in title_section.footer.paragraphs:
        for r in p.runs: r.text = ""

    if len(doc.sections) < 2:
        return

    body_section = doc.sections[1]

    # ── Header ────────────────────────────────────────────────
    header = body_section.header
    header.is_linked_to_previous = False

    # Remove existing paragraphs
    for para in list(header.paragraphs):
        para._element.getparent().remove(para._element)

    # Use _add_table_to_hf — NOT _make_table / doc.add_table
    header_table = _add_table_to_hf(header, rows=1, cols=2, total_width_emu=BODY_WIDTH)
    _remove_table_borders(header_table)
    _set_col_width(header_table, 0, BODY_WIDTH // 2)
    _set_col_width(header_table, 1, BODY_WIDTH // 2)

    lp = header_table.cell(0, 0).paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr = lp.add_run(project_name)
    lr.font.name      = T.FONT_BODY
    lr.font.size      = T.SIZE_CAPTION
    lr.font.bold      = True
    lr.font.color.rgb = T.COLOR_HEADING

    rp = header_table.cell(0, 1).paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run("TECHNICAL DOCUMENTATION")
    rr.font.name      = T.FONT_BODY
    rr.font.size      = T.SIZE_CAPTION
    rr.font.color.rgb = T.COLOR_MUTED

    # Blue rule under header
    rule_para = header.add_paragraph()
    rule_para.paragraph_format.space_before = Pt(2)
    rule_para.paragraph_format.space_after  = Pt(0)
    pPr  = rule_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm  = OxmlElement("w:bottom")
    btm.set(qn("w:val"),   "single")
    btm.set(qn("w:sz"),    "6")
    btm.set(qn("w:space"), "1")
    btm.set(qn("w:color"), "1A56DB")
    pBdr.append(btm)
    pPr.append(pBdr)

    # ── Footer ────────────────────────────────────────────────
    footer = body_section.footer
    footer.is_linked_to_previous = False

    for para in list(footer.paragraphs):
        para._element.getparent().remove(para._element)

    footer_table = _add_table_to_hf(footer, rows=1, cols=2, total_width_emu=BODY_WIDTH)
    _remove_table_borders(footer_table)
    _set_col_width(footer_table, 0, int(BODY_WIDTH * 0.6))
    _set_col_width(footer_table, 1, int(BODY_WIDTH * 0.4))

    flp = footer_table.cell(0, 0).paragraphs[0]
    flp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    flr = flp.add_run(f"© {date.today().year}  {project_name}  ·  Confidential")
    flr.font.name      = T.FONT_BODY
    flr.font.size      = T.SIZE_CAPTION
    flr.font.color.rgb = T.COLOR_MUTED

    frp = footer_table.cell(0, 1).paragraphs[0]
    frp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pre_run = frp.add_run("Page ")
    pre_run.font.name      = T.FONT_BODY
    pre_run.font.size      = T.SIZE_CAPTION
    pre_run.font.color.rgb = T.COLOR_MUTED
    _add_page_number(frp)
    frp.runs[-1].font.name      = T.FONT_BODY
    frp.runs[-1].font.size      = T.SIZE_CAPTION
    frp.runs[-1].font.color.rgb = T.COLOR_MUTED
    of_run = frp.add_run("  of  ")
    of_run.font.name      = T.FONT_BODY
    of_run.font.size      = T.SIZE_CAPTION
    of_run.font.color.rgb = T.COLOR_MUTED
    _add_numpages_field(frp)
    frp.runs[-1].font.name      = T.FONT_BODY
    frp.runs[-1].font.size      = T.SIZE_CAPTION
    frp.runs[-1].font.color.rgb = T.COLOR_MUTED


# ─────────────────────────────────────────────────────────────
# TITLE PAGE
# ─────────────────────────────────────────────────────────────

def _build_title_page(
    doc: Document,
    project_name: str,
    client_name: str,
    team_str: str,
    description: str,
):
    # ── 1. Full-width top accent bar (thick blue stripe) ──────
    #    Implemented as a table row with a tall blue-filled cell
    top_bar = _make_table(doc, rows=1, cols=1, total_width_emu=BODY_WIDTH)
    _remove_table_borders(top_bar)
    bar_cell = top_bar.cell(0, 0)
    _set_cell_shading(bar_cell, "1A56DB")

    bar_p = bar_cell.paragraphs[0]
    bar_p.paragraph_format.space_before = Pt(10)
    bar_p.paragraph_format.space_after  = Pt(10)
    bar_p.add_run("")   # empty — just for height

    # ── 2. Large vertical spacer ──────────────────────────────
    for _ in range(6):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(0)

    # ── 3. Project name (large display title) ─────────────────
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after  = Pt(4)
    title_p.paragraph_format.left_indent  = Inches(0)
    title_run = title_p.add_run(project_name.upper())
    title_run.font.name      = T.FONT_HEADING
    title_run.font.size      = T.SIZE_DISPLAY
    title_run.font.bold      = True
    title_run.font.color.rgb = T.COLOR_HEADING

    # ── 4. Subtitle label ─────────────────────────────────────
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after  = Pt(0)
    sub_run = sub_p.add_run("TECHNICAL DOCUMENTATION")
    sub_run.font.name      = T.FONT_BODY
    sub_run.font.size      = T.SIZE_SUBTITLE
    sub_run.font.color.rgb = T.COLOR_ACCENT
    sub_run.font.bold      = True
    # Letter spacing via XML
    rPr = sub_run._r.get_or_add_rPr()
    spacing_el = OxmlElement("w:spacing")
    spacing_el.set(qn("w:val"), "80")   # 8pt letter spacing (in twentieths of a pt → 160; use 60–80 for subtle)
    rPr.append(spacing_el)

    # ── 5. Accent rule ────────────────────────────────────────
    rule_p = doc.add_paragraph()
    rule_p.paragraph_format.space_before = Pt(14)
    rule_p.paragraph_format.space_after  = Pt(14)
    pPr  = rule_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm  = OxmlElement("w:bottom")
    btm.set(qn("w:val"),   "single")
    btm.set(qn("w:sz"),    "12")          # 1.5pt — noticeably heavier than before
    btm.set(qn("w:space"), "1")
    btm.set(qn("w:color"), "1A56DB")
    pBdr.append(btm)
    pPr.append(pBdr)

    # ── 6. Description ────────────────────────────────────────
    if description:
        desc_p = doc.add_paragraph()
        desc_p.paragraph_format.space_before = Pt(0)
        desc_p.paragraph_format.space_after  = Pt(28)
        desc_run = desc_p.add_run(description)
        desc_run.font.name      = T.FONT_BODY
        desc_run.font.size      = Pt(11)
        desc_run.font.color.rgb = T.COLOR_MUTED
        desc_run.font.italic    = True
    else:
        sp2 = doc.add_paragraph()
        sp2.paragraph_format.space_before = Pt(0)
        sp2.paragraph_format.space_after  = Pt(28)

    # ── 7. Metadata card ──────────────────────────────────────
    #    Two-column table: label | value
    #    Light blue background, subtle left border on entire block
    meta_rows = []
    if client_name: meta_rows.append(("Client",         client_name))
    if team_str:    meta_rows.append(("Team",           team_str))
    meta_rows.append(("Date",           date.today().strftime("%B %d, %Y")))
    meta_rows.append(("Classification", "Confidential"))

    # Outer wrapper table (single cell) for the left accent bar
    wrapper = _make_table(doc, rows=1, cols=2, total_width_emu=BODY_WIDTH)
    _remove_table_borders(wrapper)
    _set_col_width(wrapper, 0, Inches(0.06))   # narrow blue strip
    _set_col_width(wrapper, 1, BODY_WIDTH - Inches(0.06))
    _set_cell_shading(wrapper.cell(0, 0), "1A56DB")

    # The right cell contains nested table for label/value rows
    right_cell = wrapper.cell(0, 1)
    _set_cell_shading(right_cell, "F0F4FF")

    # Build the inner meta table directly inside right_cell
    # We need to add it as a nested table — use right_cell.add_table
    inner_meta = right_cell.add_table(rows=len(meta_rows), cols=2)
    inner_meta.autofit   = False
    inner_meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(inner_meta)

    # Set inner column widths
    inner_width = BODY_WIDTH - Inches(0.06)
    label_w = Inches(1.5)
    value_w = inner_width - label_w

    for row in inner_meta.rows:
        for ci, cell in enumerate(row.cells):
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW  = OxmlElement("w:tcW")
            w_val = int((label_w if ci == 0 else value_w) * 1440 / 914400)
            tcW.set(qn("w:w"),    str(w_val))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)

    for i, (label, value) in enumerate(meta_rows):
        is_last = (i == len(meta_rows) - 1)

        lc = inner_meta.cell(i, 0)
        vc = inner_meta.cell(i, 1)
        _set_cell_shading(lc, "F0F4FF")
        _set_cell_shading(vc, "F0F4FF")

        # Separator line between rows
        if not is_last:
            for cell in (lc, vc):
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBd = OxmlElement("w:tcBorders")
                btm_b = OxmlElement("w:bottom")
                btm_b.set(qn("w:val"),   "single")
                btm_b.set(qn("w:sz"),    "2")
                btm_b.set(qn("w:space"), "0")
                btm_b.set(qn("w:color"), "D0D8F0")
                tcBd.append(btm_b)
                tcPr.append(tcBd)

        lp = lc.paragraphs[0]
        lp.paragraph_format.left_indent  = Inches(0.25)
        lp.paragraph_format.space_before = Pt(7)
        lp.paragraph_format.space_after  = Pt(7)
        lr = lp.add_run(label.upper())
        lr.font.name      = T.FONT_BODY
        lr.font.size      = T.SIZE_META_LABEL
        lr.font.bold      = True
        lr.font.color.rgb = T.COLOR_LIGHT_MUTED

        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(7)
        vp.paragraph_format.space_after  = Pt(7)
        vp.paragraph_format.left_indent  = Inches(0.1)
        vr = vp.add_run(value)
        vr.font.name      = T.FONT_BODY
        vr.font.size      = T.SIZE_META_VALUE
        vr.font.color.rgb = T.COLOR_HEADING

    # ── 8. Bottom spacer then page break ─────────────────────
    for _ in range(3):
        sp3 = doc.add_paragraph()
        sp3.paragraph_format.space_before = Pt(0)
        sp3.paragraph_format.space_after  = Pt(0)

    doc.add_section(WD_SECTION.NEW_PAGE)


# ─────────────────────────────────────────────────────────────
# TABLE OF CONTENTS
# ─────────────────────────────────────────────────────────────

def _estimate_section_page_numbers(sections: List[Dict[str, Any]]) -> List[int]:
    page = 3
    start_pages = []
    for sec in sections:
        start_pages.append(page)
        word_count = len(sec.get("content", "").split())
        pages_used = max(1, round(word_count / 300))
        page += pages_used
    return start_pages


def _build_toc_page(doc: Document, sections: List[Dict[str, Any]]):
    sorted_sections = sorted(sections, key=lambda s: s.get("order", 0))
    start_pages     = _estimate_section_page_numbers(sorted_sections)

    toc_heading_p = doc.add_paragraph()
    toc_heading_p.paragraph_format.space_before = Pt(0)
    toc_heading_p.paragraph_format.space_after  = Pt(4)
    thr = toc_heading_p.add_run("Table of Contents")
    thr.font.name      = T.FONT_HEADING
    thr.font.size      = T.SIZE_H1
    thr.font.bold      = True
    thr.font.color.rgb = T.COLOR_HEADING

    _add_horizontal_rule(doc, color_hex="1A56DB", thickness="6")

    toc_table = _make_table(doc, rows=len(sorted_sections), cols=2, total_width_emu=BODY_WIDTH)
    toc_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(toc_table)
    _set_col_width(toc_table, 0, BODY_WIDTH - Inches(0.8))
    _set_col_width(toc_table, 1, Inches(0.8))

    for row_idx, (sec, start_pg) in enumerate(zip(sorted_sections, start_pages)):
        row = toc_table.rows[row_idx]
        _prevent_table_row_split(row)

        if row_idx % 2 == 0:
            for cell in row.cells:
                _set_cell_shading(cell, "F8F9FF")

        name_cell = row.cells[0]
        name_para = name_cell.paragraphs[0]
        name_para.paragraph_format.space_before = Pt(5)
        name_para.paragraph_format.space_after  = Pt(5)
        name_para.paragraph_format.left_indent  = Inches(0.12)

        section_num_run = name_para.add_run(f"{row_idx + 1}.  ")
        section_num_run.font.name      = T.FONT_BODY
        section_num_run.font.size      = T.SIZE_BODY
        section_num_run.font.bold      = True
        section_num_run.font.color.rgb = T.COLOR_ACCENT

        name_run = name_para.add_run(sec["name"])
        name_run.font.name      = T.FONT_BODY
        name_run.font.size      = T.SIZE_BODY
        name_run.font.color.rgb = T.COLOR_HEADING

        pg_cell = row.cells[1]
        pg_para = pg_cell.paragraphs[0]
        pg_para.alignment                     = WD_ALIGN_PARAGRAPH.RIGHT
        pg_para.paragraph_format.space_before = Pt(5)
        pg_para.paragraph_format.space_after  = Pt(5)
        pg_para.paragraph_format.right_indent = Inches(0.12)

        pg_run = pg_para.add_run(str(start_pg))
        pg_run.font.name      = T.FONT_BODY
        pg_run.font.size      = T.SIZE_BODY
        pg_run.font.color.rgb = T.COLOR_MUTED

    _add_horizontal_rule(doc, color_hex="D0D8F0", thickness="4")

    note_p = doc.add_paragraph()
    note_p.paragraph_format.space_before = Pt(4)
    note_p.paragraph_format.space_after  = Pt(0)
    note_run = note_p.add_run(
        "Page numbers are estimates. Exact figures may vary slightly after final rendering."
    )
    note_run.font.name      = T.FONT_BODY
    note_run.font.size      = T.SIZE_CAPTION
    note_run.font.color.rgb = T.COLOR_MUTED
    note_run.font.italic    = True

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# TABLE RENDERER
# ─────────────────────────────────────────────────────────────

def _collect_table_rows(lines: list, start_idx: int):
    rows = []
    i    = start_idx
    while i < len(lines):
        line = lines[i]
        if not (line.startswith("|") and "|" in line[1:]):
            break
        if re.match(r"^\|[\s\-|:]+\|$", line.strip()):
            i += 1
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def _render_professional_table(doc: Document, rows: List[List[str]]):
    if not rows:
        return

    col_count = max(len(r) for r in rows)
    rows      = [r + [""] * (col_count - len(r)) for r in rows]

    table           = _make_table(doc, rows=len(rows), cols=col_count, total_width_emu=BODY_WIDTH)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(table)

    for row_idx, row_data in enumerate(rows):
        is_header = (row_idx == 0)
        tr        = table.rows[row_idx]
        _prevent_table_row_split(tr)

        for col_idx, cell_text in enumerate(row_data):
            cell = tr.cells[col_idx]

            if is_header:
                _set_cell_shading(cell, "0D1B3E")
            elif row_idx % 2 == 0:
                _set_cell_shading(cell, "F8F9FF")
            else:
                _set_cell_shading(cell, "FFFFFF")

            if is_header:
                _set_cell_border(cell, {
                    "bottom": {"val": "single", "sz": "6", "color": "1A56DB"}
                })

            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(5)
            para.paragraph_format.space_after  = Pt(5)
            para.paragraph_format.left_indent  = Inches(0.08)

            if is_header:
                run = para.add_run(cell_text)
                run.font.name      = T.FONT_BODY
                run.font.size      = T.SIZE_BODY
                run.font.bold      = True
                run.font.color.rgb = T.COLOR_WHITE
            else:
                _apply_inline_markdown(para, cell_text, base_color=T.COLOR_INK)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after  = Pt(4)


# ─────────────────────────────────────────────────────────────
# INLINE MARKDOWN
# ─────────────────────────────────────────────────────────────

_INLINE_TOKEN_RE = re.compile(
    r"(\*\*\*[^*]+\*\*\*"
    r"|\*\*[^*]+\*\*"
    r"|\*[^*]+\*"
    r"|`[^`]+`"
    r"|[^*`]+"
    r"|[*`])"
)


def _apply_inline_markdown(paragraph, text: str, base_color: RGBColor = None):
    if base_color is None:
        base_color = T.COLOR_INK

    for match in _INLINE_TOKEN_RE.finditer(text):
        token = match.group(0)
        if token.startswith("***") and token.endswith("***"):
            run = paragraph.add_run(token[3:-3])
            run.bold           = True
            run.italic         = True
            run.font.name      = T.FONT_BODY
            run.font.size      = T.SIZE_BODY
            run.font.color.rgb = base_color
        elif token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold           = True
            run.font.name      = T.FONT_BODY
            run.font.size      = T.SIZE_BODY
            run.font.color.rgb = base_color
        elif token.startswith("*") and token.endswith("*") and len(token) > 1:
            run = paragraph.add_run(token[1:-1])
            run.italic         = True
            run.font.name      = T.FONT_BODY
            run.font.size      = T.SIZE_BODY
            run.font.color.rgb = base_color
        elif token.startswith("`") and token.endswith("`") and len(token) > 1:
            run = paragraph.add_run(token[1:-1])
            run.font.name      = T.FONT_MONO
            run.font.size      = T.SIZE_CODE
            run.font.color.rgb = T.COLOR_CODE_TEXT
        else:
            run = paragraph.add_run(token)
            run.font.name      = T.FONT_BODY
            run.font.size      = T.SIZE_BODY
            run.font.color.rgb = base_color


# ─────────────────────────────────────────────────────────────
# CODE BLOCK
# ─────────────────────────────────────────────────────────────

def _apply_code_paragraph_border(para):
    pPr         = para._p.get_or_add_pPr()
    pBdr        = OxmlElement("w:pBdr")
    left_border = OxmlElement("w:left")
    left_border.set(qn("w:val"),   "single")
    left_border.set(qn("w:sz"),    "12")
    left_border.set(qn("w:space"), "6")
    left_border.set(qn("w:color"), "1A56DB")
    pBdr.append(left_border)
    pPr.append(pBdr)


def _render_code_block(doc: Document, code_lines: List[str]):
    pre_spacer = doc.add_paragraph()
    pre_spacer.paragraph_format.space_before = Pt(4)
    pre_spacer.paragraph_format.space_after  = Pt(0)

    for idx, line in enumerate(code_lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.22)
        p.paragraph_format.right_indent = Inches(0.22)
        p.paragraph_format.space_before = Pt(6) if idx == 0 else Pt(0)
        p.paragraph_format.space_after  = Pt(6) if idx == len(code_lines) - 1 else Pt(0)
        _set_paragraph_shading(p, "F5F7FF")
        _apply_code_paragraph_border(p)

        run = p.add_run(line)
        run.font.name      = T.FONT_MONO
        run.font.size      = T.SIZE_CODE
        run.font.color.rgb = T.COLOR_CODE_TEXT

    post_spacer = doc.add_paragraph()
    post_spacer.paragraph_format.space_before = Pt(0)
    post_spacer.paragraph_format.space_after  = Pt(8)


# ─────────────────────────────────────────────────────────────
# HEADING DEDUPLICATION
# ─────────────────────────────────────────────────────────────

def _strip_leading_duplicate_heading(content: str, section_name: str) -> str:
    lines = content.lstrip("\n").split("\n")
    if not lines:
        return content
    first = lines[0].strip()
    if re.match(r"^#{1,3}\s+", first):
        heading_text = re.sub(r"^#{1,3}\s+", "", first).strip().lower()
        if heading_text == section_name.strip().lower():
            return "\n".join(lines[1:]).lstrip("\n")
    return content


# ─────────────────────────────────────────────────────────────
# SECTION CONTENT RENDERER
# ─────────────────────────────────────────────────────────────

def _render_section_content(doc: Document, content: str):
    lines            = content.split("\n")
    in_code_block    = False
    code_lines       = []
    i                = 0
    current_num_id   = None
    in_numbered_list = False

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code_block:
                _render_code_block(doc, code_lines)
                code_lines    = []
                in_code_block = False
            else:
                in_code_block = True
            in_numbered_list = False
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if line.startswith("|") and "|" in line[1:]:
            rows, next_i = _collect_table_rows(lines, i)
            _render_professional_table(doc, rows)
            in_numbered_list = False
            i = next_i
            continue

        is_numbered = bool(re.match(r"^\d+\.\s", line))
        if in_numbered_list and not is_numbered:
            in_numbered_list = False
            current_num_id   = None

        if line.startswith("##### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(line[6:].strip())
            r.bold           = True
            r.font.name      = T.FONT_BODY
            r.font.size      = T.SIZE_H4
            r.font.color.rgb = T.COLOR_MUTED

        elif line.startswith("#### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before   = Pt(8)
            p.paragraph_format.space_after    = Pt(2)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(line[5:].strip())
            r.bold           = True
            r.font.name      = T.FONT_BODY
            r.font.size      = T.SIZE_H4
            r.font.color.rgb = T.COLOR_HEADING

        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)

        elif line.startswith("> "):
            bq_p = doc.add_paragraph()
            bq_p.paragraph_format.left_indent  = Inches(0.3)
            bq_p.paragraph_format.space_before = Pt(4)
            bq_p.paragraph_format.space_after  = Pt(4)
            _set_paragraph_shading(bq_p, "EEF2FF")
            pPr    = bq_p._p.get_or_add_pPr()
            pBdr   = OxmlElement("w:pBdr")
            left_b = OxmlElement("w:left")
            left_b.set(qn("w:val"),   "single")
            left_b.set(qn("w:sz"),    "12")
            left_b.set(qn("w:space"), "6")
            left_b.set(qn("w:color"), "1A56DB")
            pBdr.append(left_b)
            pPr.append(pBdr)
            _apply_inline_markdown(bq_p, line[2:].strip(), base_color=T.COLOR_BLOCKQUOTE)

        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(2)
            _apply_inline_markdown(p, line[2:].strip())

        elif is_numbered:
            if not in_numbered_list:
                try:
                    current_num_id = _create_fresh_numbering(doc)
                except Exception:
                    current_num_id = None
                in_numbered_list = True
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(2)
            if current_num_id is not None:
                _apply_num_id_to_paragraph(p, current_num_id)
            _apply_inline_markdown(p, re.sub(r"^\d+\.\s", "", line).strip())

        elif re.match(r"^\*\*(.+)\*\*$", line.strip()):
            label_text = re.match(r"^\*\*(.+)\*\*$", line.strip()).group(1)
            p          = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(2)
            r          = p.add_run(label_text)
            r.bold           = True
            r.font.name      = T.FONT_BODY
            r.font.size      = Pt(11)
            r.font.color.rgb = T.COLOR_HEADING

        elif line.strip() in ("---", "***", "___"):
            _add_horizontal_rule(doc)

        elif line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = T.SPACE_AFTER_BODY
            p.paragraph_format.line_spacing = T.LINE_SPACING_BODY
            _apply_inline_markdown(p, line.strip())

        i += 1

    if in_code_block and code_lines:
        _render_code_block(doc, code_lines)


# ─────────────────────────────────────────────────────────────
# _make_table for Document bodies (title page, TOC, content)
# ─────────────────────────────────────────────────────────────

def _make_table(container, rows: int, cols: int, total_width_emu: int):
    """
    Add a table to a Document (or any container with .add_table(rows, cols)).
    For Header/Footer containers use _add_table_to_hf() instead.
    """
    table           = container.add_table(rows=rows, cols=cols)
    table.autofit   = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_width(table, total_width_emu)
    return table


# ─────────────────────────────────────────────────────────────
# PUBLIC API
# ─────────────────────────────────────────────────────────────

def build_document(
    project_id: str,
    metadata:   Dict[str, Any],
    sections:   List[Dict[str, Any]],
) -> Dict[str, Any]:
    doc = Document()
    _configure_margins(doc)
    _apply_base_styles(doc)

    project_name = metadata.get("project_name", "Untitled Project")
    client_name  = metadata.get("client_name",  "")
    team_members = metadata.get("team_members",  [])
    team_str     = ", ".join(team_members) if isinstance(team_members, list) else team_members
    description  = metadata.get("description",  "")

    _build_title_page(doc, project_name, client_name, team_str, description)
    _build_toc_page(doc, sections)

    sorted_sections = sorted(sections, key=lambda s: s.get("order", 0))

    for idx, sec in enumerate(sorted_sections):
        heading_p       = doc.add_heading(sec["name"], level=1)
        heading_p.style = doc.styles["Heading 1"]
        heading_p.paragraph_format.page_break_before = (idx > 0)
        heading_p.paragraph_format.keep_with_next    = True

        _add_horizontal_rule(doc, color_hex="D0D8F0", thickness="4")

        clean_content = _strip_leading_duplicate_heading(
            sec.get("content", ""), sec["name"]
        )
        _render_section_content(doc, clean_content)

    _configure_header_footer(doc, project_name)

    out_dir  = Path(STORAGE_DIR) / "projects" / project_id
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "output.docx"
    doc.save(str(out_path))

    full_text     = " ".join(sec.get("content", "") for sec in sections)
    word_count    = len(full_text.split())
    page_estimate = max(1, round(word_count / 300))

    return {
        "file_path":     str(out_path),
        "word_count":    word_count,
        "page_estimate": page_estimate,
        "section_count": len(sections),
    }