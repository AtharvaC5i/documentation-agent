import os
from pathlib import Path
from datetime import date
from typing import List, Dict, Any
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

load_dotenv()

STORAGE_DIR = os.getenv("STORAGE_DIR", os.path.join("..", "storage"))


# ─────────────────────────────────────────────────────────────
# DESIGN TOKENS
# ─────────────────────────────────────────────────────────────

class DesignTokens:
    COLOR_INK          = RGBColor(0x0F, 0x0F, 0x1A)
    COLOR_HEADING      = RGBColor(0x0D, 0x1B, 0x3E)
    COLOR_ACCENT       = RGBColor(0x1A, 0x56, 0xDB)
    COLOR_MUTED        = RGBColor(0x5A, 0x5A, 0x72)
    COLOR_CODE_TEXT    = RGBColor(0x1E, 0x3A, 0x8A)
    COLOR_TABLE_STRIPE = RGBColor(0xF8, 0xF9, 0xFF)
    COLOR_WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
    COLOR_BLOCKQUOTE   = RGBColor(0x2D, 0x4A, 0x8A)

    FONT_BODY    = "Calibri"
    FONT_HEADING = "Calibri"
    FONT_MONO    = "Consolas"

    SIZE_DISPLAY  = Pt(32)
    SIZE_SUBTITLE = Pt(13)
    SIZE_H1       = Pt(16)
    SIZE_H2       = Pt(13)
    SIZE_H3       = Pt(11)
    SIZE_H4       = Pt(10.5)
    SIZE_BODY     = Pt(10.5)
    SIZE_CAPTION  = Pt(9)
    SIZE_CODE     = Pt(9)

    SPACE_BEFORE_H1   = Pt(18)
    SPACE_AFTER_H1    = Pt(6)
    SPACE_BEFORE_H2   = Pt(14)
    SPACE_AFTER_H2    = Pt(4)
    SPACE_BEFORE_H3   = Pt(8)
    SPACE_AFTER_H3    = Pt(3)
    SPACE_AFTER_BODY  = Pt(6)
    LINE_SPACING_BODY = Pt(15)

    MARGIN_TOP    = Cm(2.8)
    MARGIN_BOTTOM = Cm(2.8)
    MARGIN_LEFT   = Cm(3.0)
    MARGIN_RIGHT  = Cm(2.5)

T = DesignTokens


# ─────────────────────────────────────────────────────────────
# LOW-LEVEL XML HELPERS
# ─────────────────────────────────────────────────────────────

def _set_cell_shading(cell, fill_hex: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def _set_paragraph_shading(paragraph, fill_hex: str):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    pPr.append(shd)


def _set_cell_border(cell, sides: dict):
    tc        = cell._tc
    tcPr      = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, attrs in sides.items():
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   attrs.get("val",   "single"))
        el.set(qn("w:sz"),    attrs.get("sz",    "4"))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), attrs.get("color", "000000"))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _remove_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def _prevent_table_row_split(row):
    """BUG 5 FIX: cantSplit prevents rows from breaking across pages."""
    tr   = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    cantSplit = OxmlElement("w:cantSplit")
    cantSplit.set(qn("w:val"), "1")
    trPr.append(cantSplit)


def _add_horizontal_rule(doc: Document, color_hex: str = "D0D8F0", thickness: str = "4"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm  = OxmlElement("w:bottom")
    btm.set(qn("w:val"),   "single")
    btm.set(qn("w:sz"),    thickness)
    btm.set(qn("w:space"), "1")
    btm.set(qn("w:color"), color_hex)
    pBdr.append(btm)
    pPr.append(pBdr)


def _add_page_number(paragraph):
    run      = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText      = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _add_numpages_field(paragraph):
    run      = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText      = OxmlElement("w:instrText")
    instrText.text = "NUMPAGES"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _add_toc(doc: Document):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after  = Pt(0)
    run      = paragraph.add_run()
    fldChar  = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def _insert_section_break_next_page(doc: Document):
    """BUG 7 FIX: Creates a Word section boundary so title page is isolated."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    pPr    = para._p.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")
    pgSz   = OxmlElement("w:pgSz")
    pgSz.set(qn("w:w"),      "12240")
    pgSz.set(qn("w:h"),      "15840")
    pgSz.set(qn("w:orient"), "portrait")
    sectPr.append(pgSz)
    sectPr.set(qn("w:type"), "nextPage")
    pPr.append(sectPr)


def _create_fresh_numbering(doc: Document) -> int:
    """BUG 10 FIX: New abstractNum+num per list so numbering always starts at 1."""
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        return 1

    numbering         = numbering_part._element
    abstract_nums     = numbering.findall(qn("w:abstractNum"))
    next_abstract_id  = str(len(abstract_nums))

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), next_abstract_id)

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract_num.append(multi_level)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)

    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    pPr_lvl = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"),    "360")
    ind.set(qn("w:hanging"), "360")
    pPr_lvl.append(ind)
    lvl.append(pPr_lvl)
    abstract_num.append(lvl)

    first_num = numbering.find(qn("w:num"))
    if first_num is not None:
        numbering.insert(list(numbering).index(first_num), abstract_num)
    else:
        numbering.append(abstract_num)

    existing_nums = numbering.findall(qn("w:num"))
    next_num_id   = str(len(existing_nums) + 1)

    num_el = OxmlElement("w:num")
    num_el.set(qn("w:numId"), next_num_id)

    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), next_abstract_id)
    num_el.append(abstract_ref)

    numbering.append(num_el)
    return int(next_num_id)


def _apply_num_id_to_paragraph(paragraph, num_id: int, ilvl: int = 0):
    pPr      = paragraph._p.get_or_add_pPr()
    numPr    = OxmlElement("w:numPr")
    ilvl_el  = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    existing = pPr.find(qn("w:numPr"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(numPr)


# ─────────────────────────────────────────────────────────────
# DOCUMENT-LEVEL SETUP
# ─────────────────────────────────────────────────────────────

def _configure_margins(doc: Document):
    for section in doc.sections:
        section.top_margin    = T.MARGIN_TOP
        section.bottom_margin = T.MARGIN_BOTTOM
        section.left_margin   = T.MARGIN_LEFT
        section.right_margin  = T.MARGIN_RIGHT


def _apply_base_styles(doc: Document):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name        = T.FONT_BODY
    normal.font.size        = T.SIZE_BODY
    normal.font.color.rgb   = T.COLOR_INK
    normal.paragraph_format.space_after  = T.SPACE_AFTER_BODY
    normal.paragraph_format.line_spacing = T.LINE_SPACING_BODY

    h1 = styles["Heading 1"]
    h1.font.name       = T.FONT_HEADING
    h1.font.size       = T.SIZE_H1
    h1.font.bold       = True
    h1.font.color.rgb  = T.COLOR_HEADING
    h1.font.underline  = False
    h1.paragraph_format.space_before   = T.SPACE_BEFORE_H1
    h1.paragraph_format.space_after    = T.SPACE_AFTER_H1
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name       = T.FONT_HEADING
    h2.font.size       = T.SIZE_H2
    h2.font.bold       = True
    h2.font.color.rgb  = T.COLOR_HEADING
    h2.font.underline  = False
    h2.paragraph_format.space_before   = T.SPACE_BEFORE_H2
    h2.paragraph_format.space_after    = T.SPACE_AFTER_H2
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name       = T.FONT_HEADING
    h3.font.size       = T.SIZE_H3
    h3.font.bold       = True
    h3.font.italic     = False
    h3.font.color.rgb  = T.COLOR_ACCENT
    h3.font.underline  = False
    h3.paragraph_format.space_before   = T.SPACE_BEFORE_H3
    h3.paragraph_format.space_after    = T.SPACE_AFTER_H3
    h3.paragraph_format.keep_with_next = True

    try:
        lb = styles["List Bullet"]
        lb.font.name      = T.FONT_BODY
        lb.font.size      = T.SIZE_BODY
        lb.font.color.rgb = T.COLOR_INK
        lb.paragraph_format.space_after = Pt(3)
        lb.paragraph_format.left_indent = Inches(0.25)
    except KeyError:
        pass


# ─────────────────────────────────────────────────────────────
# HEADER & FOOTER
# BUG 7 FIX: Applied only to sections[1] — title page stays clean
# ─────────────────────────────────────────────────────────────

def _configure_header_footer(doc: Document, project_name: str):
    # Clear title page section header/footer
    title_section = doc.sections[0]
    title_section.different_first_page_header_footer = False
    title_section.header.is_linked_to_previous = False
    title_section.footer.is_linked_to_previous = False
    for p in title_section.header.paragraphs:
        for r in p.runs:
            r.text = ""
    for p in title_section.footer.paragraphs:
        for r in p.runs:
            r.text = ""

    if len(doc.sections) < 2:
        return

    body_section = doc.sections[1]

    header = body_section.header
    header.is_linked_to_previous = False
    for para in header.paragraphs:
        para._element.getparent().remove(para._element)

    header_table = header.add_table(rows=1, cols=2, width=Inches(7))
    _remove_table_borders(header_table)

    lp = header_table.cell(0, 0).paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr = lp.add_run(project_name)
    lr.font.name      = T.FONT_BODY
    lr.font.size      = T.SIZE_CAPTION
    lr.font.bold      = True
    lr.font.color.rgb = T.COLOR_HEADING

    rp = header_table.cell(0, 1).paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run("TECHNICAL DOCUMENTATION")
    rr.font.name      = T.FONT_BODY
    rr.font.size      = T.SIZE_CAPTION
    rr.font.color.rgb = T.COLOR_MUTED

    rule_para = header.add_paragraph()
    rule_para.paragraph_format.space_before = Pt(2)
    rule_para.paragraph_format.space_after  = Pt(0)
    pPr  = rule_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm  = OxmlElement("w:bottom")
    btm.set(qn("w:val"),   "single")
    btm.set(qn("w:sz"),    "6")
    btm.set(qn("w:space"), "1")
    btm.set(qn("w:color"), "1A56DB")
    pBdr.append(btm)
    pPr.append(pBdr)

    footer = body_section.footer
    footer.is_linked_to_previous = False
    for para in footer.paragraphs:
        para._element.getparent().remove(para._element)

    footer_table = footer.add_table(rows=1, cols=2, width=Inches(7))
    _remove_table_borders(footer_table)

    flp = footer_table.cell(0, 0).paragraphs[0]
    flp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    flr = flp.add_run(f"© {date.today().year}  {project_name}  ·  Confidential")
    flr.font.name      = T.FONT_BODY
    flr.font.size      = T.SIZE_CAPTION
    flr.font.color.rgb = T.COLOR_MUTED

    frp = footer_table.cell(0, 1).paragraphs[0]
    frp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pre_run = frp.add_run("Page ")
    pre_run.font.name      = T.FONT_BODY
    pre_run.font.size      = T.SIZE_CAPTION
    pre_run.font.color.rgb = T.COLOR_MUTED
    _add_page_number(frp)
    frp.runs[-1].font.name      = T.FONT_BODY
    frp.runs[-1].font.size      = T.SIZE_CAPTION
    frp.runs[-1].font.color.rgb = T.COLOR_MUTED
    of_run = frp.add_run("  of  ")
    of_run.font.name      = T.FONT_BODY
    of_run.font.size      = T.SIZE_CAPTION
    of_run.font.color.rgb = T.COLOR_MUTED
    _add_numpages_field(frp)
    frp.runs[-1].font.name      = T.FONT_BODY
    frp.runs[-1].font.size      = T.SIZE_CAPTION
    frp.runs[-1].font.color.rgb = T.COLOR_MUTED


# ─────────────────────────────────────────────────────────────
# TITLE PAGE
# ─────────────────────────────────────────────────────────────

def _build_title_page(
    doc: Document,
    project_name: str,
    client_name: str,
    team_str: str,
    description: str,
):
    for _ in range(4):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after  = Pt(0)

    accent_table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(accent_table)
    accent_table.columns[0].width = Inches(0.12)
    accent_table.columns[1].width = Inches(6.5)
    _set_cell_shading(accent_table.cell(0, 0), "1A56DB")

    rp = accent_table.cell(0, 1).paragraphs[0]
    rp.paragraph_format.left_indent  = Inches(0.3)
    rp.paragraph_format.space_before = Pt(0)
    rp.paragraph_format.space_after  = Pt(0)
    title_run = rp.add_run(project_name)
    title_run.font.name      = T.FONT_HEADING
    title_run.font.size      = T.SIZE_DISPLAY
    title_run.font.bold      = True
    title_run.font.color.rgb = T.COLOR_HEADING

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(6)
    sub_p.paragraph_format.space_after  = Pt(0)
    sub_p.paragraph_format.left_indent  = Inches(0.42)
    sub_run = sub_p.add_run("TECHNICAL DOCUMENTATION")
    sub_run.font.name      = T.FONT_BODY
    sub_run.font.size      = T.SIZE_SUBTITLE
    sub_run.font.color.rgb = T.COLOR_ACCENT

    _add_horizontal_rule(doc, color_hex="1A56DB", thickness="6")

    if description:
        desc_p = doc.add_paragraph()
        desc_p.paragraph_format.left_indent  = Inches(0.42)
        desc_p.paragraph_format.space_before = Pt(6)
        desc_p.paragraph_format.space_after  = Pt(16)
        desc_run = desc_p.add_run(description)
        desc_run.font.name      = T.FONT_BODY
        desc_run.font.size      = Pt(11)
        desc_run.font.color.rgb = T.COLOR_MUTED
        desc_run.font.italic    = True

    meta_rows = []
    if client_name:
        meta_rows.append(("Client", client_name))
    if team_str:
        meta_rows.append(("Team", team_str))
    meta_rows.append(("Date",           date.today().strftime("%B %d, %Y")))
    meta_rows.append(("Classification", "Confidential"))

    meta_table = doc.add_table(rows=len(meta_rows), cols=2)
    _remove_table_borders(meta_table)
    meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (label, value) in enumerate(meta_rows):
        label_cell = meta_table.cell(i, 0)
        value_cell = meta_table.cell(i, 1)
        label_cell.width = Inches(1.6)
        value_cell.width = Inches(4.5)

        lp = label_cell.paragraphs[0]
        lp.paragraph_format.left_indent  = Inches(0.42)
        lp.paragraph_format.space_before = Pt(3)
        lp.paragraph_format.space_after  = Pt(3)
        lr = lp.add_run(label.upper())
        lr.font.name      = T.FONT_BODY
        lr.font.size      = T.SIZE_CAPTION
        lr.font.bold      = True
        lr.font.color.rgb = T.COLOR_MUTED

        vp = value_cell.paragraphs[0]
        vp.paragraph_format.space_before = Pt(3)
        vp.paragraph_format.space_after  = Pt(3)
        vr = vp.add_run(value)
        vr.font.name      = T.FONT_BODY
        vr.font.size      = T.SIZE_BODY
        vr.font.color.rgb = T.COLOR_INK

    # BUG 7 FIX: section break instead of page break
    _insert_section_break_next_page(doc)


# ─────────────────────────────────────────────────────────────
# TABLE OF CONTENTS PAGE
# BUG 8 FIX: No instruction text — clean TOC field only
# ─────────────────────────────────────────────────────────────

def _build_toc_page(doc: Document):
    toc_heading_p = doc.add_paragraph()
    toc_heading_p.paragraph_format.space_before = Pt(0)
    toc_heading_p.paragraph_format.space_after  = Pt(4)
    thr = toc_heading_p.add_run("Table of Contents")
    thr.font.name      = T.FONT_HEADING
    thr.font.size      = T.SIZE_H1
    thr.font.bold      = True
    thr.font.color.rgb = T.COLOR_HEADING

    _add_horizontal_rule(doc, color_hex="1A56DB", thickness="6")
    _add_toc(doc)
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# TABLE RENDERER
# BUG 4 FIX: inline markdown on cell content
# BUG 5 FIX: cantSplit on every row
# BUG 6 FIX: post-table spacer reduced to Pt(4)
# ─────────────────────────────────────────────────────────────

def _collect_table_rows(lines: list, start_idx: int):
    rows = []
    i    = start_idx
    while i < len(lines):
        line = lines[i]
        if not (line.startswith("|") and "|" in line[1:]):
            break
        if re.match(r"^\|[\s\-|:]+\|$", line.strip()):
            i += 1
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def _render_professional_table(doc: Document, rows: List[List[str]]):
    if not rows:
        return

    col_count = max(len(r) for r in rows)
    rows      = [r + [""] * (col_count - len(r)) for r in rows]

    table           = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(table)

    for row_idx, row_data in enumerate(rows):
        is_header = (row_idx == 0)
        tr        = table.rows[row_idx]

        # BUG 5 FIX
        _prevent_table_row_split(tr)

        for col_idx, cell_text in enumerate(row_data):
            cell = tr.cells[col_idx]

            if is_header:
                _set_cell_shading(cell, "0D1B3E")
            elif row_idx % 2 == 0:
                _set_cell_shading(cell, "F8F9FF")
            else:
                _set_cell_shading(cell, "FFFFFF")

            if is_header:
                _set_cell_border(cell, {
                    "bottom": {"val": "single", "sz": "6", "color": "1A56DB"}
                })

            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(5)
            para.paragraph_format.space_after  = Pt(5)
            para.paragraph_format.left_indent  = Inches(0.08)

            if is_header:
                run = para.add_run(cell_text)
                run.font.name      = T.FONT_BODY
                run.font.size      = T.SIZE_BODY
                run.font.bold      = True
                run.font.color.rgb = T.COLOR_WHITE
            else:
                # BUG 4 FIX: parse inline markdown in cell content
                _apply_inline_markdown(para, cell_text, base_color=T.COLOR_INK)

    # BUG 6 FIX: reduced from Pt(8)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after  = Pt(4)


# ─────────────────────────────────────────────────────────────
# INLINE MARKDOWN PARSER
# ─────────────────────────────────────────────────────────────

def _apply_inline_markdown(paragraph, text: str, base_color: RGBColor = None):
    if base_color is None:
        base_color = T.COLOR_INK

    token_pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|[^*`]+)")
    for match in token_pattern.finditer(text):
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold           = True
            run.font.name      = T.FONT_BODY
            run.font.size      = T.SIZE_BODY
            run.font.color.rgb = base_color
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic         = True
            run.font.name      = T.FONT_BODY
            run.font.size      = T.SIZE_BODY
            run.font.color.rgb = base_color
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name      = T.FONT_MONO
            run.font.size      = T.SIZE_CODE
            run.font.color.rgb = T.COLOR_CODE_TEXT
        else:
            run = paragraph.add_run(token)
            run.font.name      = T.FONT_BODY
            run.font.size      = T.SIZE_BODY
            run.font.color.rgb = base_color


# ─────────────────────────────────────────────────────────────
# CODE BLOCK RENDERER
# ─────────────────────────────────────────────────────────────

def _render_code_block(doc: Document, code_lines: List[str]):
    pre_spacer = doc.add_paragraph()
    pre_spacer.paragraph_format.space_before = Pt(4)
    pre_spacer.paragraph_format.space_after  = Pt(0)

    code_para = doc.add_paragraph()
    code_para.paragraph_format.space_before = Pt(6)
    code_para.paragraph_format.space_after  = Pt(6)
    code_para.paragraph_format.left_indent  = Inches(0.22)
    code_para.paragraph_format.right_indent = Inches(0.22)
    _set_paragraph_shading(code_para, "F5F7FF")

    for idx, line in enumerate(code_lines):
        run = code_para.add_run(line)
        run.font.name      = T.FONT_MONO
        run.font.size      = T.SIZE_CODE
        run.font.color.rgb = T.COLOR_CODE_TEXT
        if idx < len(code_lines) - 1:
            br_run = code_para.add_run()
            br     = OxmlElement("w:br")
            br_run._r.append(br)

    pPr         = code_para._p.get_or_add_pPr()
    pBdr        = OxmlElement("w:pBdr")
    left_border = OxmlElement("w:left")
    left_border.set(qn("w:val"),   "single")
    left_border.set(qn("w:sz"),    "12")
    left_border.set(qn("w:space"), "6")
    left_border.set(qn("w:color"), "1A56DB")
    pBdr.append(left_border)
    pPr.append(pBdr)

    post_spacer = doc.add_paragraph()
    post_spacer.paragraph_format.space_before = Pt(0)
    post_spacer.paragraph_format.space_after  = Pt(8)


# ─────────────────────────────────────────────────────────────
# HEADING DEDUPLICATION
# BUG 2 FIX: Strip matching leading # heading from LLM content
# ─────────────────────────────────────────────────────────────

def _strip_leading_duplicate_heading(content: str, section_name: str) -> str:
    lines = content.lstrip("\n").split("\n")
    if not lines:
        return content
    first = lines[0].strip()
    if re.match(r"^#{1,3}\s+", first):
        heading_text = re.sub(r"^#{1,3}\s+", "", first).strip().lower()
        if heading_text == section_name.strip().lower():
            return "\n".join(lines[1:]).lstrip("\n")
    return content


# ─────────────────────────────────────────────────────────────
# SECTION CONTENT RENDERER
# BUG 1 FIX: No page breaks — natural flow
# BUG 3 FIX: #### and ##### handlers
# BUG 9 FIX: > blockquote handler
# BUG 10 FIX: fresh numbering per list
# ─────────────────────────────────────────────────────────────

def _render_section_content(doc: Document, content: str):
    lines            = content.split("\n")
    in_code_block    = False
    code_lines       = []
    i                = 0
    current_num_id   = None
    in_numbered_list = False

    while i < len(lines):
        line = lines[i]

        # ── Code block ────────────────────────────────────────
        if line.strip().startswith("```"):
            if in_code_block:
                _render_code_block(doc, code_lines)
                code_lines       = []
                in_code_block    = False
            else:
                in_code_block    = True
            in_numbered_list = False
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Table ─────────────────────────────────────────────
        if line.startswith("|") and "|" in line[1:]:
            rows, next_i = _collect_table_rows(lines, i)
            _render_professional_table(doc, rows)
            in_numbered_list = False
            i = next_i
            continue

        is_numbered = bool(re.match(r"^\d+\.\s", line))
        if in_numbered_list and not is_numbered:
            in_numbered_list = False
            current_num_id   = None

        # ── BUG 3 FIX: ##### ─────────────────────────────────
        if line.startswith("##### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(line[6:].strip())
            r.bold           = True
            r.font.name      = T.FONT_BODY
            r.font.size      = T.SIZE_H4
            r.font.color.rgb = T.COLOR_MUTED

        # ── BUG 3 FIX: #### ──────────────────────────────────
        elif line.startswith("#### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before   = Pt(8)
            p.paragraph_format.space_after    = Pt(2)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(line[5:].strip())
            r.bold           = True
            r.font.name      = T.FONT_BODY
            r.font.size      = T.SIZE_H4
            r.font.color.rgb = T.COLOR_HEADING

        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)

        # ── BUG 9 FIX: Blockquote ────────────────────────────
        elif line.startswith("> "):
            bq_p = doc.add_paragraph()
            bq_p.paragraph_format.left_indent  = Inches(0.3)
            bq_p.paragraph_format.space_before = Pt(4)
            bq_p.paragraph_format.space_after  = Pt(4)
            _set_paragraph_shading(bq_p, "EEF2FF")
            pPr    = bq_p._p.get_or_add_pPr()
            pBdr   = OxmlElement("w:pBdr")
            left_b = OxmlElement("w:left")
            left_b.set(qn("w:val"),   "single")
            left_b.set(qn("w:sz"),    "12")
            left_b.set(qn("w:space"), "6")
            left_b.set(qn("w:color"), "1A56DB")
            pBdr.append(left_b)
            pPr.append(pBdr)
            _apply_inline_markdown(bq_p, line[2:].strip(), base_color=T.COLOR_BLOCKQUOTE)

        # ── Bullet list ───────────────────────────────────────
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(2)
            _apply_inline_markdown(p, line[2:].strip())

        # ── BUG 10 FIX: Numbered list with fresh numbering ───
        elif is_numbered:
            if not in_numbered_list:
                try:
                    current_num_id = _create_fresh_numbering(doc)
                except Exception:
                    current_num_id = None
                in_numbered_list = True
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(2)
            if current_num_id is not None:
                _apply_num_id_to_paragraph(p, current_num_id)
            _apply_inline_markdown(p, re.sub(r"^\d+\.\s", "", line).strip())

        # ── Standalone bold line ──────────────────────────────
        elif re.match(r"^\*\*(.+)\*\*$", line.strip()):
            label_text = re.match(r"^\*\*(.+)\*\*$", line.strip()).group(1)
            p          = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(2)
            r          = p.add_run(label_text)
            r.bold           = True
            r.font.name      = T.FONT_BODY
            r.font.size      = Pt(11)
            r.font.color.rgb = T.COLOR_HEADING

        # ── Horizontal rule ───────────────────────────────────
        elif line.strip() in ("---", "***", "___"):
            _add_horizontal_rule(doc)

        # ── Normal paragraph ──────────────────────────────────
        elif line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = T.SPACE_AFTER_BODY
            p.paragraph_format.line_spacing = T.LINE_SPACING_BODY
            _apply_inline_markdown(p, line.strip())

        i += 1

    if in_code_block and code_lines:
        _render_code_block(doc, code_lines)


# ─────────────────────────────────────────────────────────────
# PUBLIC API
# ─────────────────────────────────────────────────────────────

def build_document(
    project_id: str,
    metadata:   Dict[str, Any],
    sections:   List[Dict[str, Any]],
) -> Dict[str, Any]:
    doc = Document()
    _configure_margins(doc)
    _apply_base_styles(doc)

    project_name = metadata.get("project_name", "Untitled Project")
    client_name  = metadata.get("client_name",  "")
    team_members = metadata.get("team_members",  [])
    team_str     = ", ".join(team_members) if isinstance(team_members, list) else team_members
    description  = metadata.get("description",  "")

    _build_title_page(doc, project_name, client_name, team_str, description)
    _build_toc_page(doc)

    sorted_sections = sorted(sections, key=lambda s: s.get("order", 0))

    for idx, sec in enumerate(sorted_sections):
        # BUG 2 FIX: Heading 1 style so TOC picks it up
        heading_p = doc.add_heading(sec["name"], level=1)

        # BUG 1 FIX: page_break_before on heading instead of doc.add_page_break()
        heading_p.paragraph_format.page_break_before = (idx > 0)
        heading_p.paragraph_format.keep_with_next    = True

        _add_horizontal_rule(doc, color_hex="D0D8F0", thickness="4")

        # BUG 2 FIX: strip duplicate leading heading from LLM output
        clean_content = _strip_leading_duplicate_heading(
            sec.get("content", ""), sec["name"]
        )
        _render_section_content(doc, clean_content)
        # No page break here — content flows naturally into next section

    _configure_header_footer(doc, project_name)

    out_dir  = Path(STORAGE_DIR) / "projects" / project_id
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "output.docx"
    doc.save(str(out_path))

    full_text     = " ".join(sec.get("content", "") for sec in sections)
    word_count    = len(full_text.split())
    page_estimate = max(1, round(word_count / 300))

    return {
        "file_path":     str(out_path),
        "word_count":    word_count,
        "page_estimate": page_estimate,
        "section_count": len(sections),
    }
