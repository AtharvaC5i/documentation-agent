"""
Document Generation Pipeline — v3.
Key fixes:
- Header/footer uses tab stops NOT tables (avoids python-docx container bug)
- Explicit font.name = Calibri on every run (no font inheritance surprises)
- Smarter column width distribution per number of columns
- Smaller heading sizes (H1=13pt, H2=11pt)
- Quality note is minimal and optional
"""

import os
import re
from datetime import datetime, date
from typing import List, Dict, Any

from models.project import Project, ProjectStore
from utils.logger import info, success, divider

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ── Design tokens ────────────────────────────────────────────────────────────
class T:
    INK         = RGBColor(0x0F, 0x0F, 0x1A)
    HEADING     = RGBColor(0x0D, 0x1B, 0x3E)
    ACCENT      = RGBColor(0x1A, 0x56, 0xDB)
    MUTED       = RGBColor(0x5A, 0x5A, 0x72)
    LIGHT_MUTED = RGBColor(0x9C, 0xA3, 0xAF)
    CODE_TEXT   = RGBColor(0x1E, 0x3A, 0x8A)
    WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

    FONT  = "Calibri"
    MONO  = "Consolas"

    DISPLAY = Pt(26)
    H1      = Pt(13)
    H2      = Pt(11)
    H3      = Pt(10)
    BODY    = Pt(10)
    CAPTION = Pt(8)
    CODE    = Pt(8.5)
    META_L  = Pt(7.5)
    META_V  = Pt(10)

    MARGIN_TOP    = Cm(2.2)
    MARGIN_BOTTOM = Cm(2.2)
    MARGIN_LEFT   = Cm(2.5)
    MARGIN_RIGHT  = Cm(2.2)


PAGE_W = Cm(21)
BODY_W = PAGE_W - T.MARGIN_LEFT - T.MARGIN_RIGHT


# ── XML helpers ──────────────────────────────────────────────────────────────

def _shd(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for s in tcPr.findall(qn("w:shd")):
        tcPr.remove(s)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _no_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    bd = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        bd.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(bd)


def _set_tbl_width(table, emu: int):
    dxa = int(emu * 1440 / 914400)
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    w = tblPr.find(qn("w:tblW"))
    if w is None:
        w = OxmlElement("w:tblW")
        tblPr.append(w)
    w.set(qn("w:w"),    str(dxa))
    w.set(qn("w:type"), "dxa")


def _set_col_w(table, col: int, emu: int):
    dxa = int(emu * 1440 / 914400)
    for row in table.rows:
        tc = row.cells[col]._tc
        tcPr = tc.get_or_add_tcPr()
        w = tcPr.find(qn("w:tcW"))
        if w is None:
            w = OxmlElement("w:tcW")
            tcPr.append(w)
        w.set(qn("w:w"),    str(dxa))
        w.set(qn("w:type"), "dxa")


def _make_tbl(doc, rows: int, cols: int, width_emu: int):
    table = doc.add_table(rows=rows, cols=cols)
    table.autofit   = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_tbl_width(table, width_emu)
    return table


def _distribute_cols(table, cols: int):
    """Distribute table columns intelligently based on count."""
    ratios_map = {
        2: [0.35, 0.65],
        3: [0.14, 0.52, 0.34],
        4: [0.10, 0.44, 0.22, 0.24],
        5: [0.10, 0.35, 0.18, 0.18, 0.19],
        6: [0.09, 0.28, 0.13, 0.15, 0.18, 0.17],
    }
    ratios = ratios_map.get(cols, [1/cols] * cols)
    for ci, r in enumerate(ratios[:cols]):
        _set_col_w(table, ci, int(BODY_W * r))


def _hr(doc, color="D0D8F0", sz="4"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm  = OxmlElement("w:bottom")
    btm.set(qn("w:val"),   "single")
    btm.set(qn("w:sz"),    sz)
    btm.set(qn("w:space"), "1")
    btm.set(qn("w:color"), color)
    pBdr.append(btm)
    pPr.append(pBdr)


def _field(para, field_name: str):
    run = para.add_run()
    f1  = OxmlElement("w:fldChar");  f1.set(qn("w:fldCharType"), "begin")
    ins = OxmlElement("w:instrText"); ins.text = field_name
    f2  = OxmlElement("w:fldChar");  f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(ins); run._r.append(f2)


def _run(para, text: str, bold=False, italic=False, size: Pt = None,
         color: RGBColor = None, font: str = None) -> Any:
    """Add a run with all formatting explicitly set — no inheritance surprises."""
    r = para.add_run(text)
    r.bold    = bold
    r.italic  = italic
    r.font.name  = font or T.FONT
    r.font.size  = size or T.BODY
    r.font.color.rgb = color or T.INK
    return r


def _left_border(para, color="1A56DB", sz="12"):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    sz)
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


# ── Inline markdown ──────────────────────────────────────────────────────────

_INLINE_RE = re.compile(r"(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|[^*`]+|[*`])")


def _inline(para, text: str, color: RGBColor = None):
    color = color or T.INK
    for m in _INLINE_RE.finditer(text):
        tok = m.group(0)
        if tok.startswith("***") and tok.endswith("***") and len(tok) > 6:
            r = _run(para, tok[3:-3], bold=True, italic=True, color=color)
        elif tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
            r = _run(para, tok[2:-2], bold=True, color=color)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            r = _run(para, tok[1:-1], italic=True, color=color)
        elif tok.startswith("`") and tok.endswith("`") and len(tok) > 2:
            r = _run(para, tok[1:-1], font=T.MONO, size=T.CODE, color=T.CODE_TEXT)
        else:
            r = _run(para, tok, color=color)


# ── Document setup ────────────────────────────────────────────────────────────

def _setup(doc: Document):
    for sec in doc.sections:
        sec.top_margin    = T.MARGIN_TOP
        sec.bottom_margin = T.MARGIN_BOTTOM
        sec.left_margin   = T.MARGIN_LEFT
        sec.right_margin  = T.MARGIN_RIGHT

    n = doc.styles["Normal"]
    n.font.name = T.FONT
    n.font.size = T.BODY
    n.font.color.rgb = T.INK
    n.paragraph_format.space_after  = Pt(5)
    n.paragraph_format.line_spacing = Pt(14)

    heading_cfg = [
        ("Heading 1", T.H1, T.HEADING,  Pt(14), Pt(4)),
        ("Heading 2", T.H2, T.ACCENT,   Pt(10), Pt(3)),
        ("Heading 3", T.H3, T.MUTED,    Pt(6),  Pt(2)),
    ]
    for name, size, color, before, after in heading_cfg:
        try:
            h = doc.styles[name]
            h.font.name      = T.FONT
            h.font.size      = size
            h.font.bold      = True
            h.font.color.rgb = color
            h.font.underline = False
            h.font.italic    = False
            h.paragraph_format.space_before   = before
            h.paragraph_format.space_after    = after
            h.paragraph_format.keep_with_next = True
        except KeyError:
            pass


# ── Header / footer — tab-stop approach (reliable across all containers) ─────

def _add_right_tab(para, position_emu: int):
    """Add a right-aligned tab stop to a paragraph."""
    pPr  = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(position_emu * 1440 / 914400)))
    tabs.append(tab)
    pPr.append(tabs)


def _hf(doc: Document, project_name: str, client_name: str):
    """
    Header: project name (left) | TAB | BRD (right)
    Footer: confidential (left) | TAB | Page N of M (right)
    Uses tab stops — NOT tables. Reliable in all python-docx containers.
    """
    if len(doc.sections) < 2:
        return
    body_sec = doc.sections[1]
    body_sec.header.is_linked_to_previous = False
    body_sec.footer.is_linked_to_previous = False

    # Clear existing paragraphs
    for para in list(body_sec.header.paragraphs):
        para._element.getparent().remove(para._element)
    for para in list(body_sec.footer.paragraphs):
        para._element.getparent().remove(para._element)

    # ── Header ─────────────────────────────────────────────────────────────
    hdr_para = body_sec.header.add_paragraph()
    _add_right_tab(hdr_para, BODY_W)

    _run(hdr_para, f"{project_name}", bold=True, size=T.CAPTION, color=T.HEADING)
    hdr_para.add_run("\t").font.name = T.FONT
    _run(hdr_para, "BRD", size=T.CAPTION, color=T.MUTED)

    # Blue bottom border on header paragraph
    pPr  = hdr_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b    = OxmlElement("w:bottom")
    b.set(qn("w:val"),   "single")
    b.set(qn("w:sz"),    "4")
    b.set(qn("w:space"), "1")
    b.set(qn("w:color"), "1A56DB")
    pBdr.append(b); pPr.append(pBdr)
    hdr_para.paragraph_format.space_before = Pt(0)
    hdr_para.paragraph_format.space_after  = Pt(2)

    # ── Footer ─────────────────────────────────────────────────────────────
    ftr_para = body_sec.footer.add_paragraph()
    _add_right_tab(ftr_para, BODY_W)

    _run(ftr_para, f"© {date.today().year}  {client_name}  ·  Confidential",
         size=T.CAPTION, color=T.MUTED)
    ftr_para.add_run("\t").font.name = T.FONT

    _run(ftr_para, "Page ", size=T.CAPTION, color=T.MUTED)
    _field(ftr_para, "PAGE")
    ftr_para.runs[-1].font.name = T.FONT
    ftr_para.runs[-1].font.size = T.CAPTION
    ftr_para.runs[-1].font.color.rgb = T.MUTED
    _run(ftr_para, " of ", size=T.CAPTION, color=T.MUTED)
    _field(ftr_para, "NUMPAGES")
    ftr_para.runs[-1].font.name = T.FONT
    ftr_para.runs[-1].font.size = T.CAPTION
    ftr_para.runs[-1].font.color.rgb = T.MUTED

    ftr_para.paragraph_format.space_before = Pt(2)
    ftr_para.paragraph_format.space_after  = Pt(0)


# ── Title page ────────────────────────────────────────────────────────────────

def _title_page(doc, project_name, client_name, team_str, description, version):
    # Blue accent bar
    bar  = _make_tbl(doc, 1, 1, BODY_W)
    _no_borders(bar)
    c    = bar.cell(0, 0)
    _shd(c, "1A56DB")
    bp   = c.paragraphs[0]
    bp.paragraph_format.space_before = Pt(7)
    bp.paragraph_format.space_after  = Pt(7)
    _run(bp, "", size=T.BODY)

    for _ in range(4):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(0)

    # Project name
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after  = Pt(3)
    _run(tp, project_name.upper(), bold=True, size=T.DISPLAY, color=T.HEADING)

    # Subtitle
    sp2 = doc.add_paragraph()
    sp2.paragraph_format.space_after = Pt(0)
    _run(sp2, "BUSINESS REQUIREMENTS DOCUMENT", bold=True,
         size=Pt(10), color=T.ACCENT)

    # Rule
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after  = Pt(10)
    pPr  = rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b    = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "8")
    b.set(qn("w:space"), "1"); b.set(qn("w:color"), "1A56DB")
    pBdr.append(b); pPr.append(pBdr)

    if description:
        dp = doc.add_paragraph()
        dp.paragraph_format.space_after = Pt(20)
        _run(dp, description, italic=True, size=Pt(10), color=T.MUTED)

    # Metadata table
    meta_rows = []
    if client_name: meta_rows.append(("Client", client_name))
    if team_str:    meta_rows.append(("Prepared By", team_str))
    meta_rows.append(("Date", datetime.now().strftime("%B %d, %Y")))
    meta_rows.append(("Version", f"v{version}.0 — DRAFT"))
    meta_rows.append(("Classification", "Confidential"))

    # Wrapper with left blue strip
    wrapper = _make_tbl(doc, 1, 2, BODY_W)
    _no_borders(wrapper)
    _set_col_w(wrapper, 0, int(Inches(0.05)))
    _set_col_w(wrapper, 1, int(BODY_W - Inches(0.05)))
    _shd(wrapper.cell(0, 0), "1A56DB")
    rc = wrapper.cell(0, 1)
    _shd(rc, "F0F4FF")

    inner = rc.add_table(rows=len(meta_rows), cols=2)
    inner.autofit = False
    _no_borders(inner)
    lw = Inches(1.5)
    vw = BODY_W - Inches(0.05) - lw

    for row in inner.rows:
        for ci, cell in enumerate(row.cells):
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tw   = OxmlElement("w:tcW")
            tw.set(qn("w:w"), str(int((lw if ci == 0 else vw) * 1440 / 914400)))
            tw.set(qn("w:type"), "dxa")
            tcPr.append(tw)

    for i, (label, value) in enumerate(meta_rows):
        is_last = (i == len(meta_rows) - 1)
        lc = inner.cell(i, 0)
        vc = inner.cell(i, 1)
        _shd(lc, "F0F4FF")
        _shd(vc, "F0F4FF")

        if not is_last:
            for cell in (lc, vc):
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBd = OxmlElement("w:tcBorders")
                bb   = OxmlElement("w:bottom")
                bb.set(qn("w:val"), "single"); bb.set(qn("w:sz"), "2")
                bb.set(qn("w:space"), "0"); bb.set(qn("w:color"), "D0D8F0")
                tcBd.append(bb); tcPr.append(tcBd)

        lp = lc.paragraphs[0]
        lp.paragraph_format.left_indent  = Inches(0.15)
        lp.paragraph_format.space_before = Pt(5)
        lp.paragraph_format.space_after  = Pt(5)
        _run(lp, label.upper(), bold=True, size=T.META_L, color=T.LIGHT_MUTED)

        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(5)
        vp.paragraph_format.space_after  = Pt(5)
        vp.paragraph_format.left_indent  = Inches(0.08)
        _run(vp, value, size=T.META_V, color=T.HEADING)

    for _ in range(2):
        doc.add_paragraph()
    doc.add_section(WD_SECTION.NEW_PAGE)


# ── TOC ───────────────────────────────────────────────────────────────────────

def _estimate_pages(sections):
    page = 3
    result = []
    for s in sections:
        result.append(page)
        words = len(s.get("content", "").split())
        page += max(1, round(words / 320))
    return result


def _toc(doc, sections):
    hp = doc.add_paragraph()
    _run(hp, "Table of Contents", bold=True, size=T.H1, color=T.HEADING)
    hp.paragraph_format.space_after = Pt(3)

    _hr(doc, "1A56DB", "5")

    pages = _estimate_pages(sections)
    tbl   = _make_tbl(doc, len(sections), 2, BODY_W)
    _no_borders(tbl)
    _set_col_w(tbl, 0, int(BODY_W - Inches(0.7)))
    _set_col_w(tbl, 1, int(Inches(0.7)))

    for i, (sec, pg) in enumerate(zip(sections, pages)):
        row = tbl.rows[i]
        if i % 2 == 0:
            for cell in row.cells:
                _shd(cell, "F8F9FF")

        np = row.cells[0].paragraphs[0]
        np.paragraph_format.space_before = Pt(4)
        np.paragraph_format.space_after  = Pt(4)
        np.paragraph_format.left_indent  = Inches(0.1)
        _run(np, f"{i+1}.  ", bold=True, size=T.BODY, color=T.ACCENT)
        _run(np, sec.get("name", ""), size=T.BODY, color=T.HEADING)

        pp = row.cells[1].paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pp.paragraph_format.space_before = Pt(4)
        pp.paragraph_format.space_after  = Pt(4)
        pp.paragraph_format.right_indent = Inches(0.1)
        _run(pp, str(pg), size=T.BODY, color=T.MUTED)

    _hr(doc, "D0D8F0", "3")
    note = doc.add_paragraph()
    _run(note, "Page numbers are estimates.", italic=True, size=T.CAPTION, color=T.LIGHT_MUTED)
    doc.add_page_break()


# ── Table renderer ────────────────────────────────────────────────────────────

def _collect_table(lines, start):
    rows = []
    i    = start
    while i < len(lines):
        line = lines[i]
        if not (line.startswith("|") and "|" in line[1:]):
            break
        if re.match(r"^\|[\s\-|:]+\|$", line.strip()):
            i += 1
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def _render_table(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    rows = [r + [""] * (cols - len(r)) for r in rows]

    # Cap column count to avoid text overflow
    MAX_COLS = 7
    if cols > MAX_COLS:
        rows = [r[:MAX_COLS] for r in rows]
        cols = MAX_COLS

    tbl = _make_tbl(doc, len(rows), cols, BODY_W)
    _no_borders(tbl)
    _distribute_cols(tbl, cols)

    for ri, row_data in enumerate(rows):
        is_hdr = ri == 0
        tr     = tbl.rows[ri]

        for ci, text in enumerate(row_data[:cols]):
            cell = tr.cells[ci]
            if is_hdr:
                _shd(cell, "0D1B3E")
            elif ri % 2 == 0:
                _shd(cell, "F8F9FF")
            else:
                _shd(cell, "FFFFFF")

            if is_hdr:
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                bd   = OxmlElement("w:tcBorders")
                b    = OxmlElement("w:bottom")
                b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
                b.set(qn("w:space"), "0"); b.set(qn("w:color"), "1A56DB")
                bd.append(b); tcPr.append(bd)

            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Inches(0.06)

            # Truncate cell text that's too long
            display_text = text[:120] + "…" if len(text) > 120 else text

            if is_hdr:
                _run(p, display_text, bold=True, size=Pt(8.5), color=T.WHITE)
            else:
                _inline(p, display_text)

    doc.add_paragraph()


# ── Section content renderer ─────────────────────────────────────────────────

def _strip_dup_heading(content, section_name):
    lines = content.lstrip("\n").split("\n")
    if not lines:
        return content
    first = lines[0].strip()
    if re.match(r"^#{1,3}\s+", first):
        text = re.sub(r"^#{1,3}\s+", "", first).strip().lower()
        if text == section_name.strip().lower():
            return "\n".join(lines[1:]).lstrip("\n")
    return content


def _render_content(doc, content: str):
    lines = content.split("\n")
    i     = 0
    while i < len(lines):
        line    = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            rows, next_i = _collect_table(lines, i)
            _render_table(doc, rows)
            i = next_i
            continue

        if stripped.startswith("#####") or stripped.startswith("####"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after  = Pt(1)
            _run(p, stripped.lstrip("#").strip(), bold=True, size=T.H3, color=T.MUTED)

        elif stripped.startswith("### "):
            h = doc.add_heading(stripped[4:], level=3)
            # Ensure font is explicitly set
            for run in h.runs:
                run.font.name = T.FONT
                run.font.size = T.H3

        elif stripped.startswith("## "):
            h = doc.add_heading(stripped[3:], level=2)
            _left_border(h, "1A56DB", "12")
            for run in h.runs:
                run.font.name = T.FONT
                run.font.size = T.H2

        elif stripped.startswith("# "):
            h = doc.add_heading(stripped[2:], level=1)
            for run in h.runs:
                run.font.name = T.FONT
                run.font.size = T.H1

        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            _inline(p, stripped[2:])

        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            _inline(p, re.sub(r"^\d+\.\s", "", stripped))

        elif stripped in ("---", "***", "___"):
            _hr(doc)

        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(5)
            p.paragraph_format.line_spacing = Pt(14)
            _inline(p, stripped)

        i += 1


# ── Output path ───────────────────────────────────────────────────────────────

def _get_output_path(project_name: str, version: int) -> str:
    out_dir = os.path.join(os.path.dirname(__file__), "../../outputs")
    os.makedirs(out_dir, exist_ok=True)
    safe = re.sub(r'[^\w\-]', '_', project_name)
    return os.path.join(out_dir, f"BRD_{safe}_v{version}.docx")


# ── Main pipeline ─────────────────────────────────────────────────────────────

class DocumentPipeline:

    def __init__(self, project: Project, store: ProjectStore):
        self.project = project
        self.store   = store

    async def run(self):
        self.project.status           = "generating_document"
        self.project.progress_message = "Generating Word document..."
        self.store.save(self.project)
        divider(f"DOCUMENT PIPELINE — {self.project.project_name}")

        if not DOCX_AVAILABLE:
            self.project.status           = "error"
            self.project.progress_message = "python-docx not installed."
            self.store.save(self.project)
            return

        doc = Document()
        _setup(doc)

        team_str = ", ".join(self.project.team_members) if self.project.team_members else "Project Team"

        info("DOC", "Building title page...")
        _title_page(
            doc,
            self.project.project_name,
            self.project.client_name,
            team_str,
            self.project.description,
            self.project.version,
        )

        # Only approved sections; fallback to all if none approved
        approved = [s for s in self.project.generated_sections if s.get("approved")]
        if not approved:
            approved = self.project.generated_sections

        info("DOC", f"Building TOC ({len(approved)} sections)...")
        _toc(doc, approved)

        info("DOC", f"Rendering {len(approved)} sections...")
        for idx, sec in enumerate(approved):
            heading = doc.add_heading(sec.get("name", "Section"), level=1)
            if idx > 0:
                heading.paragraph_format.page_break_before = True
            heading.paragraph_format.keep_with_next = True
            _left_border(heading, "D0D8F0", "3")
            for run in heading.runs:
                run.font.name = T.FONT
                run.font.size = T.H1
            _hr(doc, "D0D8F0", "3")

            content = _strip_dup_heading(sec.get("content", ""), sec.get("name", ""))
            _render_content(doc, content)
            info("DOC", f"  [{idx+1}/{len(approved)}] {sec.get('name','')} — {sec.get('word_count', '?')} words")

        info("DOC", "Adding header/footer...")
        _hf(doc, self.project.project_name, self.project.client_name)

        output_path = _get_output_path(self.project.project_name, self.project.version)
        doc.save(output_path)

        self.project.document_path    = output_path
        self.project.status           = "complete"
        self.project.progress_message = "Document ready for download."
        self.store.save(self.project)
        success("DOC", f"Saved → {output_path}")
