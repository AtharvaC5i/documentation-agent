"""
Requirement Traceability Matrix.

Builds a matrix linking every BRD section to the requirements
it was generated from, traced back to the original source
(transcript timestamp or user story ID).

No LLM call — pure data assembly from existing project data.
Outputs both in-memory matrix and an Excel-like .docx table.
"""

import os
from typing import List, Dict, Any
from datetime import datetime

from models.project import Project, ProjectStore
from utils.logger import info, success, divider
from utils.file_utils import get_output_path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def build_traceability_matrix(project: Project) -> List[Dict[str, Any]]:
    """
    Build the traceability matrix from existing project data.
    No LLM call — uses source_location and req_id already tracked.

    Returns a list of rows:
    {
      section_name, req_id, req_type, description,
      source, source_location, speaker, priority
    }
    """
    divider("TRACEABILITY MATRIX — Building")

    req_by_id = {r.req_id: r for r in project.requirements_pool}
    matrix_rows = []

    for section in project.generated_sections:
        section_name = section.get("name", "Unknown")
        source_req_ids = section.get("source_req_ids", [])

        if not source_req_ids:
            # Section has no explicit source — add a placeholder row
            matrix_rows.append({
                "section_name": section_name,
                "req_id": "—",
                "req_type": "—",
                "description": "No specific requirements traced to this section",
                "source": "—",
                "source_location": "—",
                "speaker": "—",
                "priority": "—",
                "quality_score": section.get("quality_pct", "—"),
            })
            continue

        for req_id in source_req_ids:
            req = req_by_id.get(req_id)
            if not req:
                continue
            matrix_rows.append({
                "section_name": section_name,
                "req_id": req.req_id,
                "req_type": req.type.replace("_", " ").title(),
                "description": req.description[:200],
                "source": req.source.replace("_", " ").title(),
                "source_location": req.source_location or req.user_story_id or "—",
                "speaker": req.speaker or "—",
                "priority": (req.priority or "—").replace("_", " ").title(),
                "quality_score": section.get("quality_pct", "—"),
            })

    info("TRACEABILITY", f"Matrix built: {len(matrix_rows)} rows across {len(project.generated_sections)} sections")
    return matrix_rows


def generate_traceability_document(project: Project, matrix: List[Dict]) -> str:
    """
    Generate a Word document containing the traceability matrix.
    Returns the file path.
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed")

    safe_name = project.project_name.replace(" ", "_").replace("/", "_")
    output_dir = os.path.join(os.path.dirname(__file__), "../../outputs")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"Traceability_{safe_name}_v{project.version}.docx")

    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2)

    # ── Cover ──────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("REQUIREMENT TRACEABILITY MATRIX")
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"{project.project_name}  |  {project.client_name}  |  v{project.version}  |  {datetime.now().strftime('%B %d, %Y')}")
    meta.runs[0].font.size = Pt(11)
    meta.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    # ── Summary stats ──────────────────────────────────────────────────────
    total_reqs = len(set(r["req_id"] for r in matrix if r["req_id"] != "—"))
    total_sections = len(set(r["section_name"] for r in matrix))
    transcript_reqs = sum(1 for r in matrix if r["source"] == "Transcript")
    story_reqs = sum(1 for r in matrix if r["source"] == "User Story")

    stats = doc.add_paragraph()
    stats.add_run(
        f"Total Requirements Traced: {total_reqs}   |   "
        f"BRD Sections: {total_sections}   |   "
        f"From Transcript: {transcript_reqs}   |   "
        f"From User Stories: {story_reqs}"
    ).font.size = Pt(10)

    doc.add_paragraph()

    # ── Matrix table ───────────────────────────────────────────────────────
    headers = ["BRD Section", "Req ID", "Type", "Requirement Description", "Source", "Location / Reference", "Speaker", "Priority"]
    col_widths = [Cm(3.5), Cm(1.8), Cm(2.2), Cm(6.5), Cm(2.0), Cm(3.0), Cm(2.5), Cm(2.0)]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr.cells[i]
        cell.width = w
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Calibri"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Navy background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F3A5F")
        tcPr.append(shd)

    # Data rows with alternating shading
    prev_section = None
    row_idx = 0
    for row_data in matrix:
        row = table.add_row()
        vals = [
            row_data["section_name"],
            row_data["req_id"],
            row_data["req_type"],
            row_data["description"],
            row_data["source"],
            row_data["source_location"],
            row_data["speaker"],
            row_data["priority"],
        ]
        fill = "EFF6FF" if row_idx % 2 == 0 else "FFFFFF"
        for i, (val, w) in enumerate(zip(vals, col_widths)):
            cell = row.cells[i]
            cell.width = w
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(8.5)
            run.font.name = "Calibri"
            if i == 0 and val != prev_section:
                run.font.bold = True
                prev_section = val
            # Row shading
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
        row_idx += 1

    doc.add_paragraph()

    # ── Notes ──────────────────────────────────────────────────────────────
    note = doc.add_paragraph()
    note.add_run("Notes: ").font.bold = True
    note.add_run(
        "This matrix was auto-generated by BRD Agent. "
        "Each row traces a BRD section to the requirement that informed it, "
        "and back to the original source (meeting transcript or user story). "
        "Timestamps refer to the meeting recording. US-XXX IDs refer to Jira/user story exports."
    ).font.size = Pt(9)
    note.runs[-1].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.save(output_path)
    success("TRACEABILITY", f"Document saved: {output_path}")
    return output_path
