"""
file_extractor.py

Extracts clean text from PDF, DOCX, and TXT/MD files.
All extraction happens server-side for maximum accuracy.
"""

import io
import pdfplumber
from docx import Document
from typing import Tuple


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def extract_text(filename: str, file_bytes: bytes) -> Tuple[str, str]:
    """
    Extract text from uploaded file bytes.

    Returns:
        (extracted_text, warning_message)
        warning_message is empty string if no issues.
    """
    ext = _get_extension(filename)

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: '{ext}'. "
            f"Supported types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if ext == ".pdf":
        return _extract_pdf(file_bytes)
    elif ext == ".docx":
        return _extract_docx(file_bytes)
    else:
        return _extract_text(file_bytes)


def _get_extension(filename: str) -> str:
    lower = filename.lower()
    for ext in SUPPORTED_EXTENSIONS:
        if lower.endswith(ext):
            return ext
    # fallback: take everything after last dot
    if "." in filename:
        return "." + filename.rsplit(".", 1)[-1].lower()
    return ""


def _extract_pdf(file_bytes: bytes) -> Tuple[str, str]:
    """
    Extract text from a digital PDF using pdfplumber.
    Preserves table structure by converting tables to markdown-style grids.
    """
    warning = ""
    sections = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        total_pages = len(pdf.pages)

        for page_num, page in enumerate(pdf.pages, start=1):
            page_sections = []

            # --- Extract tables first (higher fidelity than raw text for structured data) ---
            tables = page.extract_tables()
            table_bboxes = []

            for table in tables:
                if not table:
                    continue
                # Get bounding box of this table to exclude from plain text extraction
                table_settings = {
                    "vertical_strategy": "lines",
                    "horizontal_strategy": "lines",
                }
                # Convert table rows to markdown-style text
                rows = []
                for row in table:
                    cleaned_row = [
                        (cell or "").strip().replace("\n", " ") for cell in row
                    ]
                    rows.append(" | ".join(cleaned_row))
                if rows:
                    page_sections.append("\n".join(rows))

            # --- Extract remaining text (excluding table areas) ---
            # Crop away table regions for cleaner text extraction
            cropped = page
            for t in page.find_tables():
                try:
                    bbox = t.bbox  # (x0, top, x1, bottom)
                    cropped = cropped.outside_bbox(bbox)
                except Exception:
                    pass

            text = cropped.extract_text(
                x_tolerance=3,
                y_tolerance=3,
                layout=True,           # preserve spatial layout
                x_density=7.25,
                y_density=13,
            )

            if text and text.strip():
                page_sections.append(text.strip())

            if page_sections:
                sections.append(
                    f"--- Page {page_num} of {total_pages} ---\n"
                    + "\n\n".join(page_sections)
                )

    if not sections:
        warning = (
            "No text could be extracted. "
            "This may be a scanned/image-based PDF. "
            "Please copy-paste the text manually instead."
        )
        return "", warning

    return "\n\n".join(sections), warning


def _extract_docx(file_bytes: bytes) -> Tuple[str, str]:
    """
    Extract text from a Word .docx file using python-docx.
    Preserves heading hierarchy and table structure.
    """
    warning = ""
    sections = []

    doc = Document(io.BytesIO(file_bytes))

    # --- Heading + paragraph extraction ---
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""

        if "Heading 1" in style_name:
            sections.append(f"\n# {text}")
        elif "Heading 2" in style_name:
            sections.append(f"\n## {text}")
        elif "Heading 3" in style_name:
            sections.append(f"\n### {text}")
        else:
            sections.append(text)

    # --- Table extraction ---
    for table_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            # Deduplicate merged cells (python-docx repeats merged cell content)
            deduped = []
            prev = None
            for cell in cells:
                if cell != prev:
                    deduped.append(cell)
                prev = cell
            rows.append(" | ".join(deduped))
        if rows:
            sections.append(f"\n[Table {table_idx + 1}]\n" + "\n".join(rows))

    if not sections:
        warning = "The Word document appears to be empty or contains only images."
        return "", warning

    return "\n".join(sections), warning


def _extract_text(file_bytes: bytes) -> Tuple[str, str]:
    """Extract plain text from .txt or .md files."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(encoding), ""
        except UnicodeDecodeError:
            continue
    # Last resort
    return file_bytes.decode("utf-8", errors="replace"), "Some characters could not be decoded and were replaced."